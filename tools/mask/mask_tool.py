#!/usr/bin/env python3
"""
マスクツール: 取扱説明書等から客先名・設備番号等の機密情報を隠蔽する

対応形式:
  PDF       → 黒塗り（PyMuPDF による真のリダクション）
  DOCX      → ラベル置換（例: [客先名]）
  TXT/MD/その他 → ラベル置換

インストール:
  pip install -r requirements.txt

使い方:
  # 単一ファイル
  python mask_tool.py input.pdf --config mask_config.yaml
  python mask_tool.py input.docx --config mask_config.yaml

  # 出力先を明示
  python mask_tool.py input.pdf --config mask_config.yaml --output masked.pdf

  # フォルダ内の全ファイルを一括処理
  python mask_tool.py ./manual_draft/ --config mask_config.yaml

  # 出力サフィックスを変更（デフォルト: _masked）
  python mask_tool.py input.docx --config mask_config.yaml --suffix _公開版
"""

import argparse
import re
import shutil
import sys
import unicodedata
from pathlib import Path
from typing import NamedTuple

try:
    import yaml
except ImportError:
    sys.exit("エラー: pip install pyyaml が必要です")


class MaskRule(NamedTuple):
    pattern: re.Pattern
    label: str


# ハイフン・ダッシュ類（全角/半角・長音記号を含む）。表記ゆれ許容パターンと
# 監査用の文字正規化(_norm_char)の両方で同じ集合を使う。
HYPHEN_CHARS = "‐‑‒–—―−ー－-"
_HYPHEN_ALT = "|".join(re.escape(ch) for ch in dict.fromkeys(HYPHEN_CHARS))
_DIGIT_CLASS = "[0-9０-９]"
# 全角/半角スペース（「姓　名」のように区切りに使われる。全角スペースは通常の
# 空白文字とは別扱いされるため、単純な\sだけでは拾えない）
SPACE_CHARS = " 　"
_SPACE_CLASS = "[ 　]"


def tolerant_pattern(value: str) -> str:
    """値の完全一致(re.escape)の代わりに、全角/半角の数字・ハイフン類・スペースの
    表記ゆれを許容する正規表現パターン文字列を作る。

    PDFの黒塗り(_redact_text_occurrences)は文字単位で正規化して照合するため
    表記ゆれに強いが、DOCX/XLSX/テキストのラベル置換(apply_rules)は単純な
    文字列一致のため、同じ電話番号でもセルによって全角/半角が混在していると
    一部だけマスクされ一部が素通りする、という不具合につながっていた。

    さらに、日本語の帳票では字幅を揃えるために「遠　藤」「遠 藤」のように
    文字と文字の間へ空白を挿入することがある(均等割り付けの手打ち版)。
    見た目も検索結果も同じ値なのに、素の「遠藤」だけがマスクされて
    空白入りが残るため、文字の間に空白があっても一致するようにする。
    """
    parts = []
    for c in value:
        nc = unicodedata.normalize("NFKC", c)
        if len(nc) == 1 and nc.isdigit():
            parts.append(_DIGIT_CLASS)
        elif c in HYPHEN_CHARS:
            parts.append(f"(?:{_HYPHEN_ALT})")
        elif c in SPACE_CHARS:
            parts.append(f"{_SPACE_CLASS}+")
        else:
            parts.append(re.escape(c))
    # 文字の間に挿入されうる空白を許容する(値自体が空白を含む位置は上で処理済み)
    return f"{_SPACE_CLASS}*".join(parts)


def load_rules(config_path: str) -> list[MaskRule]:
    with open(config_path, encoding="utf-8") as f:
        data = yaml.safe_load(f)

    rules = []
    for m in data.get("masks", []):
        label = m.get("label", "[MASKED]")
        if "pattern" in m:
            try:
                rules.append(MaskRule(pattern=re.compile(m["pattern"]), label=label))
            except re.error as e:
                print(f"警告: 正規表現エラー ({m['pattern']}): {e} — スキップします")
        elif "value" in m:
            rules.append(MaskRule(pattern=re.compile(tolerant_pattern(m["value"])), label=label))
        else:
            print(f"警告: 'value' または 'pattern' がないエントリをスキップ: {m}")
    return rules


def _decode_text_bytes(raw: bytes) -> str:
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("cp932", errors="replace")


def apply_rules(text: str, rules: list[MaskRule]) -> tuple[str, int]:
    count = 0
    for rule in rules:
        new_text, n = rule.pattern.subn(rule.label, text)
        count += n
        text = new_text
    return text, count


# ---------- テキストファイル ----------

def mask_text_file(input_path: Path, output_path: Path, rules: list[MaskRule]) -> int:
    text = _decode_text_bytes(input_path.read_bytes())
    new_text, count = apply_rules(text, rules)
    output_path.write_text(new_text, encoding="utf-8")
    return count


# ---------- PDF ----------

def extract_pdf_images(pdf_bytes: bytes) -> list[dict]:
    """PDF内の画像を抽出してbase64サムネイル付きで返す"""
    try:
        import fitz
    except ImportError:
        return []
    try:
        from PIL import Image as PilImage
        import io, base64
    except ImportError:
        return []

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    seen: set[int] = set()
    result = []

    for page in doc:
        for img in page.get_images():
            xref = img[0]
            if xref in seen:
                continue
            seen.add(xref)
            try:
                img_data = doc.extract_image(xref)
                pil = PilImage.open(io.BytesIO(img_data["image"]))
                pil.thumbnail((120, 120))
                buf = io.BytesIO()
                pil.save(buf, format="PNG")
                b64 = base64.b64encode(buf.getvalue()).decode()
                result.append({
                    "xref": xref,
                    "width": img_data["width"],
                    "height": img_data["height"],
                    "thumbnail": f"data:image/png;base64,{b64}",
                })
            except Exception:
                pass

    doc.close()
    return result


def extract_page_text_candidates(pdf_bytes: bytes, page_num: int = 0) -> list[dict]:
    """指定ページのテキストを行単位で抽出し、座標付きで返す（重複除去・短行除外）"""
    try:
        import fitz
    except ImportError:
        return []

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    if page_num >= len(doc):
        doc.close()
        return []

    page = doc[page_num]
    data = page.get_text("dict")
    doc.close()

    seen: set[str] = set()
    results: list[dict] = []

    for block in data.get("blocks", []):
        if block.get("type") != 0:  # テキストブロック以外はスキップ
            continue
        for line in block.get("lines", []):
            text = "".join(span["text"] for span in line.get("spans", [])).strip()
            if len(text) < 2 or text in seen:
                continue
            seen.add(text)
            bbox = line["bbox"]  # (x0, y0, x1, y1)
            results.append({
                "text": text,
                "bbox": list(bbox),
                "page": page_num,
            })

    return results


# 主要な日本人の姓（個人名の推定候補検出用。網羅ではなく代表的なもののみ）
_COMMON_SURNAMES = [
    "佐藤", "鈴木", "高橋", "田中", "渡辺", "伊藤", "山本", "中村", "小林", "加藤",
    "吉田", "山田", "佐々木", "山口", "松本", "井上", "木村", "林", "斉藤", "斎藤",
    "清水", "山崎", "森", "池田", "橋本", "阿部", "石川", "前田", "藤田", "小川",
    "岡田", "後藤", "長谷川", "村上", "近藤", "石井", "坂本", "遠藤", "青木", "藤井",
    "西村", "福田", "太田", "三浦", "岡本", "松田", "中川", "中野", "原田", "小野",
    "田村", "竹内", "金子", "和田", "中山", "石田", "上田", "森田", "柴田", "菅原",
    "酒井", "工藤", "横山", "宮崎", "宮本", "内田", "高木", "安藤", "谷口", "大野",
    "丸山", "今井", "高田", "藤本", "武田", "上野", "杉山", "増田", "小島", "村田",
    "平野", "千葉", "岩崎", "松尾", "菊地", "野口", "松井", "梅田", "山下", "新井",
    "杉本", "野村", "小山", "水野", "浜田", "北村", "秋山", "大塚", "西田", "服部",
    "小泉", "小池", "谷", "望月", "河野", "荒木", "本田", "落合", "浅野", "松岡",
    "植田", "永井", "大西", "久保", "小松", "岸本", "宮城", "斎木", "須藤", "堀",
    "篠原", "熊谷", "宇野", "牧野", "上原", "石橋", "村田", "白石", "高野", "河合",
    "北川", "森本", "田口", "大橋", "三宅", "村松", "馬場", "岸田", "森下", "白井",
    "西山", "堀内", "戸田", "岩本", "植村", "板垣", "小森", "沢田", "宮田", "村岡",
    "重松", "深田", "南", "北野", "西川", "東", "桜井", "柳沢", "柳原", "柴崎",
    "菅野", "菅井", "須田", "笹川", "笹本", "笠原", "小笠原", "小早川", "小山田", "小関",
    "神谷", "神田", "神戸", "神山", "石塚", "石黒", "石渡", "岩田", "岩下", "岩井",
    "秋元", "秋田", "安田", "安達", "阿部", "青山", "青柳", "赤松", "浅井", "浅田",
    "荒井", "有馬", "生田", "五十嵐", "池上", "池谷", "石岡", "泉", "磯部", "市川",
    "一色", "伊東", "稲垣", "稲葉", "井田", "井関", "岩瀬", "岩橋", "植木", "上村",
    "宇都宮", "梅原", "梅本", "遠山", "大石", "大川", "大久保", "大倉", "大坂", "大澤",
    "大島", "大谷", "大槻", "大平", "大森", "大山", "岡崎", "岡野", "岡部", "小笠",
    "奥山", "奥村", "尾崎", "尾形", "尾上", "角田", "梶原", "梶山", "片岡", "片桐",
    "勝又", "勝田", "加瀬", "門脇", "金井", "金沢", "金田", "金森", "上條", "唐沢",
    "河原", "河村", "菊池", "岸", "北条", "喜多", "木下", "北原", "北山", "木下",
    "楠", "熊倉", "久米", "栗原", "栗田", "小出", "小柳", "小澤", "駒沢", "小林",
    "斎木", "境", "阪本", "崎山", "笹原", "佐野", "佐久間", "沢井", "重田", "篠田",
    "柴山", "島田", "島崎", "清川", "須賀", "杉江", "杉浦", "杉田", "須永", "関口",
    "妹尾", "曽根", "高垣", "高城", "高瀬", "高柳", "高山", "滝沢", "武井", "竹下",
    "田島", "田代", "田沢", "田端", "田畑", "田辺", "田淵", "玉井", "玉川", "千田",
    "築山", "土屋", "堤", "手塚", "寺尾", "寺田", "寺島", "戸川", "土肥", "冨田",
    "冨永", "豊田", "永田", "長沼", "永野", "中井", "中尾", "永岡", "永島", "中嶋",
    "中西", "永瀬", "中谷", "中津", "中根", "中原", "中丸", "中森", "永山", "夏目",
    "難波", "西尾", "西岡", "西田", "西野", "西堀", "西脇", "布施", "沼田", "根本",
    "野沢", "野田", "野中", "野々村", "萩原", "橋口", "橋詰", "長谷部", "畑中", "服部",
    "浜口", "浜崎", "早川", "林田", "原口", "原島", "半田", "東海林", "樋口", "肥田",
    "菱田", "平井", "平田", "平沼", "深沢", "福井", "福岡", "福島", "福原", "藤井",
    "藤岡", "藤木", "藤崎", "藤沢", "藤野", "藤原", "船橋", "古川", "古田", "細川",
    "細野", "堀田", "堀口", "本間", "前川", "前野", "牧田", "町田", "松浦", "松崎",
    "松永", "松原", "松山", "丸田", "三上", "水谷", "溝口", "三田", "湊", "南野",
    "宮沢", "宮下", "宮武", "向井", "村山", "元木", "森岡", "森谷", "森山", "門田",
    "矢野", "柳", "山内", "山岡", "山川", "山北", "山越", "山根", "山之内", "山室",
    "山室", "湯浅", "横田", "横川", "吉川", "吉岡", "吉沢", "吉村", "米田", "渡部",
]
# 重複除去 + 長い姓を優先してマッチさせる（例: 「山之内」が「山内」等の部分一致で切れないように）
_SURNAMES_UNIQUE = sorted(dict.fromkeys(_COMMON_SURNAMES), key=len, reverse=True)
_SURNAME_ALT = "|".join(_SURNAMES_UNIQUE)
# 姓のみで記載されたケースの検出用。1文字姓（森・林・東・南・谷 等）は普通名詞や
# 地名と紛れて誤検出が多すぎるため、2文字以上の姓に限定する。
_SURNAME_MULTICHAR_ALT = "|".join(s for s in _SURNAMES_UNIQUE if len(s) >= 2)

# 会社名の構成文字（漢字・カタカナ・半角全角英数字・記号のみ。\w は日本語の助詞(ひらがな)も
# 「単語文字」とみなすため使わない。含めると「〇〇株式会社が受注した」の「が受注した」まで
# 巻き込んで一致してしまう）
_NAME_CHARS = r"一-龥ァ-ヶーA-Za-z0-9・\-－～（）()"

# 敬称（様・御中）の直前から名前として遡ってよい文字。ひらがなを含めないことで
# 「〜について〇〇様」の助詞部分で止まり、名前だけを取り出せる。
_HONORIFIC_TARGET_CHARS = r"一-龥ァ-ヶーA-Za-z0-9＆&・"
# 「仕様」「同様」「多様」「模様」「一様」等、様が付く一般語の直前文字。
# これらで終わる文字列は人名・会社名ではないため候補にしない。
_SAMA_EXCLUDE_TAIL = r"仕同多模一異神王"

# 氏名の直後に付く役職・敬称。これらを名前の一部として取り込むと
# 「遠藤課長」のような複合語が候補になり、チェックしてもその表記の箇所しか
# 消えず、姓だけで書かれた大多数のセルが残ってしまう。名前の切れ目として扱う。
_TITLE_SUFFIXES = [
    "課長代理", "部長代理", "工場長", "取締役", "代表", "社長", "専務", "常務",
    "本部長", "事業部長", "部長", "次長", "課長", "係長", "室長", "所長", "店長",
    "主任", "主査", "主幹", "班長", "組長", "監督", "技師", "技士", "博士",
    "先生", "殿", "様", "氏", "君", "さん", "担当", "御中",
]
_TITLE_ALT = "|".join(sorted(_TITLE_SUFFIXES, key=len, reverse=True))

# スキャン用パターン（会社名・住所は新規、その他は既存パターンから値を拾う）
# 末尾に (?P<v>...) のような明示グループがある場合はその部分のみ、なければ全体一致を候補値とする
SCAN_PATTERNS = [
    ("会社名", rf"(?:株式会社|有限会社|合同会社|合資会社|合名会社"
               rf"|一般社団法人|公益社団法人|一般財団法人|公益財団法人|特定非営利活動法人|独立行政法人)"
               rf"[{_NAME_CHARS}]{{1,30}}"),
    ("会社名", rf"[{_NAME_CHARS}]{{2,30}}(?:株式会社|有限会社|合同会社)"),
    # (株)(有) 等の略記（前株・後株どちらも。全角/半角カッコと㈱㈲の合字ライゲーションに対応）
    ("会社名(略記)", rf"(?:\(株\)|（株）|㈱|\(有\)|（有）|㈲|\(同\)|（同）)[{_NAME_CHARS}]{{1,30}}"),
    ("会社名(略記)", rf"[{_NAME_CHARS}]{{2,30}}(?:\(株\)|（株）|㈱|\(有\)|（有）|㈲|\(同\)|（同）)"),
    ("会社名", r"\b[A-Z][A-Za-z0-9&.,\-' ]{1,40}(?:Corporation|Corp\.?|Incorporated|Inc\.?|Co\.,?\s*Ltd\.?|LLC|K\.K\.)"),
    # 法人格が付かない社名（例: 日東工業・〇〇電機・〇〇製作所）。日本の製造業では
    # 取引先を法人格抜きで書くことが多く、上の株式会社系パターンでは拾えない。
    # 直前がひらがなの場合は文章の途中なので除外する（「板金工業」等の誤検出を抑える）。
    ("会社名(業種語尾)",
     rf"(?<![ぁ-ん]) [{_HONORIFIC_TARGET_CHARS}]{{2,12}}"
     rf"(?:工業|電機|電気|製作所|工作所|鉄工所|精機|機械|産業|技研|工機|製鋼|化成|商会|商事|樹脂|金属)"
     .replace(" ", "")),
    ("住所",   r"(?:北海道|東京都|大阪府|京都府|神奈川県|埼玉県|千葉県|愛知県|福岡県|兵庫県"
               r"|静岡県|茨城県|広島県|宮城県|新潟県|長野県|岐阜県|栃木県|群馬県|岡山県"
               r"|三重県|熊本県|鹿児島県|山口県|愛媛県|長崎県|奈良県|青森県|岩手県|大分県"
               r"|石川県|山形県|富山県|秋田県|香川県|和歌山県|山梨県|福島県|徳島県|高知県"
               r"|島根県|宮崎県|鳥取県|福井県|佐賀県|沖縄県)"
               r"[\S]{1,60}(?:市|区|町|村)[\S]{1,40}"),
    ("電話番号", r"0\d{1,4}[\-－]\d{1,4}[\-－]\d{4}"),
    ("携帯番号", r"0[789]0[\-－]\d{4}[\-－]\d{4}"),
    ("FAX番号",  r"FAX[：:\s]*0\d{1,4}[\-－]\d{1,4}[\-－]\d{4}"),
    ("メール",   r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"),
    ("郵便番号", r"〒?\d{3}[\-－]\d{4}"),
    ("IPアドレス", r"\b(?:192\.168|10\.\d{1,3}|172\.(?:1[6-9]|2\d|3[01]))\.\d{1,3}\.\d{1,3}\b"),

    # ---- ここから: 会社名/プロジェクト名/個人名など「形式のない自由記述」の推定候補 ----
    # ラベル一致（〇〇：△△ の△△部分を候補にする。誤検出が少なく精度が高い）
    ("会社名(ラベル)",     r"(?:客先名|お客様名|得意先名|貴社名|社名|発注者名?)[：:]\s*([^\n\r、。]{2,40})"),
    ("プロジェクト名(ラベル)", r"(?:件名|案件名|プロジェクト名|工事名|物件名|設備名)[：:]\s*([^\n\r、。]{2,60})"),
    # 担当者名は氏名しか入らないため、名前らしい文字だけに限定する
    # （[^、。]で貪欲に取ると「担当者名：変更内容」の形式で…」のような説明文を丸ごと拾う）
    ("担当者名(ラベル)",   rf"(?:担当者?名?|作成者|お客様担当|ご担当)[：:][ 　]*([{_HONORIFIC_TARGET_CHARS}]{{2,12}}(?:[ 　][{_HONORIFIC_TARGET_CHARS}]{{1,8}})?)"),

    # 敬称一致（会社名・個人名どちらにも付く）。
    # 敬称の直前から「名前らしい文字」だけを遡って取る。文中の助詞等(ひらがな)で
    # 止めないと「タッチパネルについて〇〇様」から『タッチパネルについて〇〇』を
    # 拾ってしまい、その1箇所しかマスクされない候補ができてしまう。
    # また「仕様」「同様」「様々」等の一般語を人名と誤認しないよう前後を除外する。
    ("会社名/個人名(御中)", rf"([{_HONORIFIC_TARGET_CHARS}]{{2,30}})[ 　]*御中"),
    ("会社名/個人名(様)",
     rf"([{_HONORIFIC_TARGET_CHARS}]{{2,15}})(?<![{_SAMA_EXCLUDE_TAIL}])[ 　]*様(?![々式子方邸相])"),

    # 姓の辞書一致による個人名推定（名は漢字/カタカナのみ対象とし、助詞等のひらがなを含めない。
    # 姓と名の間の全角/半角スペース区切り「鈴木　一郎」「鈴木 一郎」にも対応。誤検出あり・要確認）
    # 名の部分に役職・敬称(課長/様/氏 等)を取り込まないこと。取り込むと「遠藤課長」が
    # 候補になり、これは完全一致で消すため、姓だけで書かれた他のセルが全て残ってしまう。
    # 末尾も、続くのが役職・敬称ならそこで名前が終わったとみなす
    # （「鈴木一郎課長」から「鈴木一郎」を取り出すため）
    # 名の各文字位置で役職・敬称の始まりでないことを確認する。末尾だけ見ると
    # 「鈴木一郎様」の『様』まで名前に取り込んでしまうため。
    ("個人名(推定)",
     rf"(?<![一-龥々])(?:{_SURNAME_ALT})[ 　]?(?:(?!(?:{_TITLE_ALT}))[一-龥ァ-ヶー]){{1,3}}"
     rf"(?:(?=[ 　]?(?:{_TITLE_ALT}))|(?![一-龥ァ-ヶ]))"),

    # 姓のみの記載（管理台帳の担当者欄など、名を伴わず姓だけ書かれるケース）。
    # 上の「姓+名」パターンは名が最低1文字必要なため、姓だけのセルは候補に出ない。
    # 名が続く場合は上のパターンに任せるが、続くのが役職・敬称(課長/様 等)なら
    # それは名前ではないので、ここで姓だけを候補にする。
    # 1文字姓（森・林・東・南・谷 等）は普通名詞と紛れて誤検出が多すぎるため対象外。
    ("姓のみ(推定)",
     rf"(?<![一-龥々])(?:{_SURNAME_MULTICHAR_ALT})(?:(?=[ 　]?(?:{_TITLE_ALT}))|(?![ 　]?[一-龥ァ-ヶー]))"),
]


def scan_text_candidates(text: str, extra_patterns: list[dict] | None = None) -> list[dict]:
    """任意のテキストからマスク候補を検出して返す（PDF/DOCX/テキスト共通）

    パターンに捕捉グループ(...)がある場合はグループ1の値のみを候補にする
    （「客先名：〇〇」のようなラベル一致で、ラベル部分を候補に含めないため）。
    """
    patterns = list(SCAN_PATTERNS)
    for ep in (extra_patterns or []):
        patterns.append((ep.get("label", "その他"), ep["pattern"]))

    seen: set[str] = set()
    results: list[dict] = []
    for label, pat in patterns:
        try:
            for m in re.finditer(pat, text):
                val = (m.group(1) if m.lastindex else m.group(0)).strip()
                # マスク済みの置換ラベル(《客先名》等)を候補として出さない。
                # 一度マスクしたファイルを再度かけたときに無意味な候補が並ぶため。
                if "《" in val or "》" in val:
                    continue
                if val and val not in seen:
                    seen.add(val)
                    results.append({"value": val, "label": label})
        except re.error:
            pass

    return results


def scan_pdf_candidates(pdf_bytes: bytes, extra_patterns: list[dict] | None = None) -> list[dict]:
    """PDFテキストからマスク候補を検出して返す"""
    try:
        import fitz
    except ImportError:
        return []

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    full_text = "\n".join(page.get_text() for page in doc)
    doc.close()
    return scan_text_candidates(full_text, extra_patterns)


def scan_docx_candidates(source, extra_patterns: list[dict] | None = None) -> list[dict]:
    """DOCXテキストからマスク候補を検出して返す（sourceはパスまたはfile-like）"""
    return scan_text_candidates(extract_docx_text(source), extra_patterns)


def scan_plain_text_candidates(raw: bytes, extra_patterns: list[dict] | None = None) -> list[dict]:
    """テキストファイルの内容からマスク候補を検出して返す"""
    return scan_text_candidates(_decode_text_bytes(raw), extra_patterns)


def scan_file_candidates(filename: str, raw: bytes, extra_patterns: list[dict] | None = None) -> list[dict]:
    """ファイル種別を判定してマスク候補を検出して返す（PDF/DOCX/テキスト共通の入口）"""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return scan_pdf_candidates(raw, extra_patterns)
    elif suffix == ".docx":
        import io
        return scan_docx_candidates(io.BytesIO(raw), extra_patterns)
    elif suffix == ".xlsx":
        import io
        return scan_xlsx_candidates(io.BytesIO(raw), extra_patterns)
    elif suffix in TEXT_SUFFIXES or suffix == "":
        return scan_plain_text_candidates(raw, extra_patterns)
    return []


def render_pdf_pages(pdf_bytes: bytes, max_pages: int = 10, scale: float = 1.5, start_page: int = 0) -> list[dict]:
    """PDF各ページをPNG画像としてレンダリングして返す"""
    try:
        import fitz
        import base64
    except ImportError:
        return []

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    mat = fitz.Matrix(scale, scale)
    pages = []

    end = min(start_page + max_pages, len(doc))
    for i in range(start_page, end):
        page = doc[i]
        pix = page.get_pixmap(matrix=mat)
        b64 = base64.b64encode(pix.tobytes("png")).decode()
        pages.append({
            "page": i,
            "pdf_width": page.rect.width,
            "pdf_height": page.rect.height,
            "image": f"data:image/png;base64,{b64}",
        })

    doc.close()
    return pages


def _norm_char(c: str) -> str:
    """照合用の文字正規化: 全角/半角の差・ハイフン類の字種差を吸収する"""
    c = unicodedata.normalize("NFKC", c)
    # ハイフン・ダッシュ類はすべて '-' に寄せる
    return "".join("-" if ch in HYPHEN_CHARS else ch for ch in c)


def _generalize_pattern(tn: str) -> str:
    """正規化済み文字列を形式パターン(正規表現)に汎化する。
    例: '5263-1649F' → r'\\d\\d\\d\\d-\\d\\d\\d\\d[A-Za-z]'
    数字→\\d、半角英字→[A-Za-z]、その他(記号・日本語)はリテラルのまま。
    """
    parts = []
    for c in tn:
        if c.isdigit():
            parts.append(r"\d")
        elif c.isalpha() and c.isascii():
            parts.append("[A-Za-z]")
        else:
            parts.append(re.escape(c))
    return "".join(parts)


def _norm_str(s: str) -> str:
    return "".join(_norm_char(c) for c in s if not c.isspace())


def _build_audit_matchers(literal_texts: list[str] | None, generalize: bool) -> tuple[list[str], list[re.Pattern]]:
    """監査(残存チェック)用の照合対象を組み立てる。戻り値: (正規化済みリテラル一覧, 汎化パターン一覧)"""
    literals = [_norm_str(t) for t in (literal_texts or []) if _norm_str(t)]
    gen_pats = []
    if generalize:
        for tn in literals:
            if len(tn) >= 6:
                try:
                    gen_pats.append(re.compile(_generalize_pattern(tn)))
                except re.error:
                    pass
    return literals, gen_pats


def audit_pdf(pdf_path: Path, literal_texts: list[str] | None = None,
              regex_rules: list[MaskRule] | None = None,
              generalize: bool = False) -> list[dict]:
    """マスク済みPDFを再スキャンし、対象文字列・形式・パターンが残っているページを返す。

    数百ページのPDFでも、目視確認が必要なページだけを特定できる。
    戻り値: [{"page": ページ番号(1始まり), "hits": [検出内容, ...]}, ...]
    """
    try:
        import fitz
    except ImportError:
        return []

    literals, gen_pats = _build_audit_matchers(literal_texts, generalize)

    doc = fitz.open(str(pdf_path))
    flagged = []
    for i, page in enumerate(doc):
        raw_text = page.get_text()
        stream = _norm_str(raw_text)
        hits = []
        for lit in literals:
            if lit in stream:
                hits.append(f"文字列: {lit[:20]}")
        for pat in gen_pats:
            m = pat.search(stream)
            if m:
                hits.append(f"形式一致: {m.group()[:20]}")
        for rule in (regex_rules or []):
            m = rule.pattern.search(raw_text)
            if m:
                hits.append(f"{rule.label}: {m.group()[:20]}")
        if hits:
            flagged.append({"page": i + 1, "hits": sorted(set(hits))})
    doc.close()
    return flagged


def audit_text_blob(text: str, literal_texts: list[str] | None = None,
                     regex_rules: list[MaskRule] | None = None,
                     generalize: bool = False, unit: str = "行") -> list[dict]:
    """任意のテキストを行単位で再スキャンし、対象文字列・形式・パターンが残っている行を返す。

    DOCX/テキストファイルの監査に使う（PDFのページ単位に相当する行/段落単位）。
    戻り値: [{"location": "行N" 等, "hits": [検出内容, ...]}, ...]
    """
    literals, gen_pats = _build_audit_matchers(literal_texts, generalize)

    flagged = []
    for i, line in enumerate(text.splitlines()):
        if not line.strip():
            continue
        stream = _norm_str(line)
        hits = []
        for lit in literals:
            if lit in stream:
                hits.append(f"文字列: {lit[:20]}")
        for pat in gen_pats:
            m = pat.search(stream)
            if m:
                hits.append(f"形式一致: {m.group()[:20]}")
        for rule in (regex_rules or []):
            m = rule.pattern.search(line)
            if m:
                hits.append(f"{rule.label}: {m.group()[:20]}")
        if hits:
            flagged.append({"location": f"{unit}{i + 1}", "hits": sorted(set(hits))})
    return flagged


def audit_docx(source, literal_texts: list[str] | None = None,
               regex_rules: list[MaskRule] | None = None,
               generalize: bool = False) -> list[dict]:
    """マスク済みDOCXを再スキャンし、対象文字列・形式・パターンが残っている段落を返す"""
    return audit_text_blob(extract_docx_text(source), literal_texts, regex_rules, generalize, unit="段落")


def audit_text_file(raw: bytes, literal_texts: list[str] | None = None,
                     regex_rules: list[MaskRule] | None = None,
                     generalize: bool = False) -> list[dict]:
    """マスク済みテキストファイルを再スキャンし、対象文字列・形式・パターンが残っている行を返す"""
    return audit_text_blob(_decode_text_bytes(raw), literal_texts, regex_rules, generalize, unit="行")


def audit_file(filename: str, raw: bytes, literal_texts: list[str] | None = None,
               regex_rules: list[MaskRule] | None = None,
               generalize: bool = False) -> list[dict]:
    """ファイル種別を判定してマスク後の残存チェックを行う（PDFを除く共通の入口。PDFはaudit_pdfを直接使用）"""
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        import io
        return audit_docx(io.BytesIO(raw), literal_texts, regex_rules, generalize)
    elif suffix == ".xlsx":
        import io
        return audit_xlsx(io.BytesIO(raw), literal_texts, regex_rules, generalize)
    elif suffix in TEXT_SUFFIXES or suffix == "":
        return audit_text_file(raw, literal_texts, regex_rules, generalize)
    return []


def _redact_text_occurrences(page, targets: list[str], generalize: bool = False) -> int:
    """ページ内の全文字を空白無視で連結し、対象文字列に一致する箇所を黒塗りする。

    - テキストが複数のオブジェクトに分割されていても(例: '5'+'ABC-123'+'C')全体をカバー
    - 全角/半角・ハイフン字種の違いを正規化して照合
    - 一致部分の前後にスペースなしで密着している文字(枝番・改訂記号など)も塗り広げる
    - 一致文字がその行の過半を占める場合は行全体を黒塗り
      (図面欄のように1文字ずつ間隔をあけて配置され、シート番号・改訂記号が
       スペースを挟んで同じ欄に入っているケースをカバーする)
    - generalize=True の場合、対象文字列を形式パターンに汎化して照合する
      (例: '5263-1649F' → 数字4桁-数字4桁+英字1字 の形式すべてに一致。
       誤爆防止のため6文字以上の候補のみ汎化し、短い候補はリテラル一致)
    """
    import fitz

    def norm_str(s: str) -> str:
        return "".join(_norm_char(c) for c in s if not c.isspace())

    target_norms = [norm_str(t) for t in targets if norm_str(t)]
    if not target_norms:
        return 0

    # 照合パターンを構築(リテラル or 汎化形式)
    patterns: list[re.Pattern] = []
    for tn in target_norms:
        if generalize and len(tn) >= 6:
            patterns.append(re.compile(_generalize_pattern(tn)))
        else:
            patterns.append(re.compile(re.escape(tn)))

    # sort=True で視覚的な並び順(上→下、左→右)に文字を取得する
    # スペースは「区切りマーカー」として保持する(塗り広げの停止条件に使う)
    items: list[tuple[str, "fitz.Rect | None", int]] = []  # (char, rect, line_id)
    line_bboxes: list["fitz.Rect"] = []
    line_char_counts: list[int] = []  # 行ごとの非スペース文字数
    for block in page.get_text("rawdict", sort=True).get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            line_id = len(line_bboxes)
            line_bboxes.append(fitz.Rect(line["bbox"]))
            n_chars = 0
            for span in line.get("spans", []):
                for ch in span.get("chars", []):
                    c = ch.get("c", "")
                    if not c:
                        continue
                    if c.isspace():
                        items.append((" ", None, line_id))
                    else:
                        items.append((c, fitz.Rect(ch["bbox"]), line_id))
                        n_chars += 1
            line_char_counts.append(n_chars)
            items.append((" ", None, line_id))  # 行末も区切り扱い

    # 正規化ストリームと、各位置→itemsインデックスの対応表を作る
    stream_chars: list[str] = []
    stream_map: list[int] = []
    for i, (c, rect, _) in enumerate(items):
        if rect is None:
            continue
        for nc in _norm_char(c):
            stream_chars.append(nc)
            stream_map.append(i)
    stream = "".join(stream_chars)

    def _adjacent(a, b):  # aがbの左隣か
        h = max(a.height, b.height)
        same_row = abs((a.y0 + a.y1) - (b.y0 + b.y1)) / 2 < h * 0.7
        gap = b.x0 - a.x1
        return same_row and -h * 0.5 < gap < h * 0.35

    count = 0
    for pat in patterns:
        for m in pat.finditer(stream):
            idx = m.start()
            mlen = m.end() - m.start()
            if mlen <= 0:
                continue
            lo = stream_map[idx]
            hi = stream_map[idx + mlen - 1]

            # 前後に密着している文字まで塗り広げる
            # (スペースで停止。テキストオブジェクトの境界は幾何的な隣接判定で越える)
            while lo - 1 >= 0 and items[lo - 1][1] is not None and _adjacent(items[lo - 1][1], items[lo][1]):
                lo -= 1
            while hi + 1 < len(items) and items[hi + 1][1] is not None and _adjacent(items[hi][1], items[hi + 1][1]):
                hi += 1

            # 行ごとの一致文字数を数える
            matched_per_line: dict[int, int] = {}
            for i in range(lo, hi + 1):
                if items[i][1] is not None:
                    matched_per_line[items[i][2]] = matched_per_line.get(items[i][2], 0) + 1

            covered_lines: set[int] = set()
            char_rects: list["fitz.Rect"] = []
            for line_id, n_matched in matched_per_line.items():
                if line_char_counts[line_id] > 0 and n_matched / line_char_counts[line_id] >= 0.5:
                    # 行の過半が一致 → 行全体(欄全体)を黒塗り
                    covered_lines.add(line_id)
                else:
                    # 文中の一部 → 一致した文字だけ黒塗り
                    char_rects.extend(items[i][1] for i in range(lo, hi + 1)
                                      if items[i][1] is not None and items[i][2] == line_id)

            for line_id in covered_lines:
                page.add_redact_annot(line_bboxes[line_id] + (-1, -1, 1, 1), fill=(0, 0, 0))
                count += 1

            if char_rects:
                cur = fitz.Rect(char_rects[0])
                for r in char_rects[1:]:
                    same_row = abs((r.y0 + r.y1) - (cur.y0 + cur.y1)) / 2 < max(r.height, cur.height) * 0.7
                    if same_row:
                        cur |= r
                    else:
                        page.add_redact_annot(cur + (-1, -1, 1, 1), fill=(0, 0, 0))
                        count += 1
                        cur = fitz.Rect(r)
                page.add_redact_annot(cur + (-1, -1, 1, 1), fill=(0, 0, 0))
                count += 1

    return count


def mask_pdf(
    input_path: Path,
    output_path: Path,
    rules: list[MaskRule],
    image_xrefs: list[int] | None = None,
    regions: list[dict] | None = None,
    line_texts: list[str] | None = None,
    generalize: bool = False,
) -> int:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        sys.exit("エラー: pip install pymupdf が必要です")

    doc = fitz.open(str(input_path))
    total_pages = len(doc)
    total = 0
    xref_set = set(image_xrefs or [])

    # 範囲指定黒塗り: allPages=true の場合は全ページに展開
    # ページサイズが指定元と異なる場合(A4/A3混在等)は比率でスケーリングする
    region_map: dict[int, list] = {}
    for r in (regions or []):
        targets = range(total_pages) if r.get("allPages") else [int(r["page"])]
        src_w = float(r.get("pdfW", 0)) or None
        src_h = float(r.get("pdfH", 0)) or None
        for pg in targets:
            pw, ph = doc[pg].rect.width, doc[pg].rect.height
            sx = pw / src_w if src_w else 1.0
            sy = ph / src_h if src_h else 1.0
            rect = fitz.Rect(r["x0"] * sx, r["y0"] * sy, r["x1"] * sx, r["y1"] * sy)
            region_map.setdefault(pg, []).append(rect)

    for i, page in enumerate(doc):
        # テキスト黒塗り
        page_text = page.get_text()
        matched_strings: set[str] = set()
        for rule in rules:
            for m in rule.pattern.finditer(page_text):
                matched_strings.add(m.group())
        for s in matched_strings:
            for rect in page.search_for(s):
                page.add_redact_annot(rect, fill=(0, 0, 0))
                total += 1

        # 画像黒塗り
        for img in page.get_images():
            xref = img[0]
            if xref in xref_set:
                try:
                    bbox = page.get_image_bbox(img)
                    page.add_redact_annot(bbox, fill=(0, 0, 0))
                    total += 1
                except Exception:
                    pass

        # 範囲指定黒塗り
        for rect in region_map.get(i, []):
            page.add_redact_annot(rect, fill=(0, 0, 0))
            total += 1

        # テキスト候補: 文字単位で照合して黒塗り
        # （番号が複数のテキストオブジェクトに分割されている場合や
        #   スペース有無の違いがあっても、文字列全体をカバーする）
        if line_texts:
            total += _redact_text_occurrences(page, line_texts, generalize)

        # 画像・ベクターグラフィックにも黒塗りを適用
        try:
            page.apply_redactions(
                images=fitz.PDF_REDACT_IMAGE_PIXELS,
                graphics=fitz.PDF_REDACT_LINE_ART_REMOVE_IF_COVERED,
            )
        except TypeError:
            page.apply_redactions()

    doc.save(str(output_path), garbage=4, deflate=True)
    doc.close()
    return total


# ---------- DOCX ----------

def extract_docx_text(source) -> str:
    """DOCXから本文・表・ヘッダー/フッターの全テキストを抽出する（sourceはパスまたはfile-like）"""
    try:
        from docx import Document
    except ImportError:
        return ""

    doc = Document(source)
    parts: list[str] = []

    def collect_paragraphs(paragraphs):
        for p in paragraphs:
            if p.text.strip():
                parts.append(p.text)

    def collect_table(table):
        for row in table.rows:
            for cell in row.cells:
                collect_paragraphs(cell.paragraphs)
                for nested in cell.tables:
                    collect_table(nested)

    collect_paragraphs(doc.paragraphs)
    for table in doc.tables:
        collect_table(table)

    for section in doc.sections:
        collect_paragraphs(section.header.paragraphs)
        collect_paragraphs(section.footer.paragraphs)
        for hdr in (section.even_page_header, section.first_page_header):
            collect_paragraphs(hdr.paragraphs)
        for ftr in (section.even_page_footer, section.first_page_footer):
            collect_paragraphs(ftr.paragraphs)

    return "\n".join(parts)


def _process_paragraphs(paragraphs, rules: list[MaskRule]) -> int:
    count = 0
    for para in paragraphs:
        full = para.text
        if not full.strip():
            continue
        new_full, n = apply_rules(full, rules)
        if n == 0:
            continue
        count += n
        runs = para.runs
        if not runs:
            continue
        # runs を先頭にまとめて置換（段落内インライン書式は先頭runの書式を引き継ぐ）
        runs[0].text = new_full
        for run in runs[1:]:
            run.text = ""
    return count


def mask_docx(input_path: Path, output_path: Path, rules: list[MaskRule]) -> int:
    try:
        from docx import Document
    except ImportError:
        sys.exit("エラー: pip install python-docx が必要です")

    shutil.copy2(input_path, output_path)
    doc = Document(str(output_path))
    total = 0

    # 本文
    total += _process_paragraphs(doc.paragraphs, rules)

    # テーブル（ネスト対応）
    def process_table(table):
        nonlocal total
        for row in table.rows:
            for cell in row.cells:
                total += _process_paragraphs(cell.paragraphs, rules)
                for nested in cell.tables:
                    process_table(nested)

    for table in doc.tables:
        process_table(table)

    # ヘッダー・フッター
    for section in doc.sections:
        total += _process_paragraphs(section.header.paragraphs, rules)
        total += _process_paragraphs(section.footer.paragraphs, rules)
        # 偶数・奇数ページのヘッダー/フッター
        for hdr in (section.even_page_header, section.first_page_header):
            total += _process_paragraphs(hdr.paragraphs, rules)
        for ftr in (section.even_page_footer, section.first_page_footer):
            total += _process_paragraphs(ftr.paragraphs, rules)

    doc.save(str(output_path))
    return total


# ---------- XLSX ----------

def _open_xlsx(source, data_only: bool = False):
    """XLSXを開く（sourceはパス、またはfile-like。file-likeは複数回開けるようseek(0)する）"""
    import openpyxl
    if hasattr(source, "seek"):
        source.seek(0)
    return openpyxl.load_workbook(source, data_only=data_only)


def _is_formula_cell(cell) -> bool:
    val = cell.value
    return cell.data_type == "f" or (isinstance(val, str) and val.startswith("="))


# 印刷ヘッダー/フッターは「奇数/偶数/先頭ページ」×「左/中央/右」で保持される
_HF_GROUPS = ("oddHeader", "oddFooter", "evenHeader", "evenFooter", "firstHeader", "firstFooter")
_HF_PARTS = ("left", "center", "right")


def _iter_xlsx_header_footer(ws):
    """シートの印刷ヘッダー/フッターの文字列を (場所の説明, 文字列) で列挙する"""
    for group_name in _HF_GROUPS:
        group = getattr(ws, group_name, None)
        if group is None:
            continue
        for part_name in _HF_PARTS:
            part = getattr(group, part_name, None)
            text = getattr(part, "text", None)
            if isinstance(text, str) and text.strip():
                yield f"{group_name}.{part_name}", text


def _iter_xlsx_all_texts(source):
    """全シートの「文字列として扱えるセル値」を (シート名, セル位置, 表示テキスト, 数式か) で列挙する。

    通常の文字列セルに加え、数式セル（例: =B2&" "&C2 で氏名を結合しているセル）についても
    Excelが保存時に書き込むキャッシュ済み計算結果(data_only=True)を対象にする。数式セルは
    値を書き換えると数式が壊れるため、検出・監査の対象としては拾うが、実際の値の変更は
    mask_xlsx側でルールに一致したセルのみ行う（一致しない数式は温存する）。
    キャッシュが存在しない（一度もExcelで計算・保存されていない）場合は検出できない。
    """
    try:
        wb = _open_xlsx(source, data_only=False)
        wb_cached = _open_xlsx(source, data_only=True)
    except ImportError:
        return
    try:
        for ws in wb.worksheets:
            ws_cached = wb_cached[ws.title] if ws.title in wb_cached.sheetnames else None
            for row in ws.iter_rows():
                for cell in row:
                    if _is_formula_cell(cell):
                        if ws_cached is None:
                            continue
                        cached_val = ws_cached[cell.coordinate].value
                        if isinstance(cached_val, str) and cached_val.strip():
                            yield ws.title, cell.coordinate, cached_val, True
                    else:
                        val = cell.value
                        if isinstance(val, str) and val.strip():
                            yield ws.title, cell.coordinate, val, False
                    # セルのメモ（コメント）。セル値とは別に保持されるため
                    # セルだけ見ていると氏名が残る。
                    comment = getattr(cell, "comment", None)
                    if comment is not None and isinstance(comment.text, str) and comment.text.strip():
                        yield ws.title, f"{cell.coordinate}(メモ)", comment.text, False
            # 印刷時のヘッダー/フッター。画面上のセルには現れないが文字列として保存される。
            for where, text in _iter_xlsx_header_footer(ws):
                yield ws.title, where, text, False
    finally:
        wb.close()
        wb_cached.close()


def extract_xlsx_text(source) -> str:
    """XLSXの全シートの文字列（数式セルはキャッシュ済み計算結果を含む）を改行区切りで抽出する
    （sourceはパスまたはfile-like）"""
    return "\n".join(val for _, _, val, _ in _iter_xlsx_all_texts(source))


def scan_xlsx_candidates(source, extra_patterns: list[dict] | None = None) -> list[dict]:
    """XLSXの文字列セルからマスク候補を検出して返す"""
    return scan_text_candidates(extract_xlsx_text(source), extra_patterns)


def mask_xlsx(input_path: Path, output_path: Path, rules: list[MaskRule]) -> int:
    try:
        import openpyxl
    except ImportError:
        sys.exit("エラー: pip install openpyxl が必要です")

    shutil.copy2(input_path, output_path)
    wb = openpyxl.load_workbook(str(output_path), data_only=False)
    wb_cached = openpyxl.load_workbook(str(output_path), data_only=True)
    total = 0

    for ws in wb.worksheets:
        ws_cached = wb_cached[ws.title] if ws.title in wb_cached.sheetnames else None
        for row in ws.iter_rows():
            for cell in row:
                if _is_formula_cell(cell):
                    # 数式セル: キャッシュ済み計算結果がルールに一致する場合のみ、
                    # マスク後の文字列で数式ごと置き換える（機密情報を数式経由で
                    # 再現させないことを、数式を保つことより優先する）
                    if ws_cached is None:
                        continue
                    cached_val = ws_cached[cell.coordinate].value
                    if not isinstance(cached_val, str) or not cached_val:
                        continue
                    new_val, n = apply_rules(cached_val, rules)
                    if n:
                        cell.value = new_val
                        total += n
                else:
                    val = cell.value
                    if isinstance(val, str) and val:
                        new_val, n = apply_rules(val, rules)
                        if n:
                            cell.value = new_val
                            total += n

                # セルのメモ（コメント）はセル値とは別に保存されるため個別に処理する
                comment = getattr(cell, "comment", None)
                if comment is not None and isinstance(comment.text, str) and comment.text:
                    new_text, n = apply_rules(comment.text, rules)
                    if n:
                        comment.text = new_text
                        total += n

        # 印刷ヘッダー/フッター（画面のセルには出ないが文字列として保存される）
        for group_name in _HF_GROUPS:
            group = getattr(ws, group_name, None)
            if group is None:
                continue
            for part_name in _HF_PARTS:
                part = getattr(group, part_name, None)
                text = getattr(part, "text", None)
                if isinstance(text, str) and text:
                    new_text, n = apply_rules(text, rules)
                    if n:
                        part.text = new_text
                        total += n

    wb.save(str(output_path))
    wb.close()
    wb_cached.close()
    return total


def audit_xlsx(source, literal_texts: list[str] | None = None,
               regex_rules: list[MaskRule] | None = None,
               generalize: bool = False) -> list[dict]:
    """マスク済みXLSXを再スキャンし、対象文字列・形式・パターンが残っているセルを返す。

    戻り値: [{"location": "シート名!セル位置" (数式セルは末尾に(数式)), "hits": [検出内容, ...]}, ...]
    """
    literals, gen_pats = _build_audit_matchers(literal_texts, generalize)

    flagged = []
    for sheet_name, coord, raw_val, is_formula in _iter_xlsx_all_texts(source):
        stream = _norm_str(raw_val)
        hits = []
        for lit in literals:
            if lit in stream:
                hits.append(f"文字列: {lit[:20]}")
        for pat in gen_pats:
            m = pat.search(stream)
            if m:
                hits.append(f"形式一致: {m.group()[:20]}")
        for rule in (regex_rules or []):
            m = rule.pattern.search(raw_val)
            if m:
                hits.append(f"{rule.label}: {m.group()[:20]}")
        if hits:
            loc = f"{sheet_name}!{coord}" + ("(数式)" if is_formula else "")
            flagged.append({"location": loc, "hits": sorted(set(hits))})
    return flagged


# ---------- ディスパッチ ----------

TEXT_SUFFIXES = {".txt", ".md", ".csv", ".log", ".ini", ".json", ".xml", ".html", ".htm"}

def process_file(input_path: Path, output_path: Path, rules: list[MaskRule], image_xrefs: list[int] | None = None, regions: list[dict] | None = None, line_texts: list[str] | None = None, generalize: bool = False):
    suffix = input_path.suffix.lower()
    if suffix == ".pdf":
        count = mask_pdf(input_path, output_path, rules, image_xrefs, regions, line_texts, generalize)
        print(f"  [PDF ]  {count:4d}箇所 黒塗り  → {output_path.name}")
    elif suffix == ".docx":
        count = mask_docx(input_path, output_path, rules)
        print(f"  [DOCX]  {count:4d}箇所 ラベル置換 → {output_path.name}")
    elif suffix == ".xlsx":
        count = mask_xlsx(input_path, output_path, rules)
        print(f"  [XLSX]  {count:4d}箇所 ラベル置換 → {output_path.name}")
    elif suffix in TEXT_SUFFIXES or suffix == "":
        count = mask_text_file(input_path, output_path, rules)
        print(f"  [TEXT]  {count:4d}箇所 ラベル置換 → {output_path.name}")
    else:
        # 非対応形式はコピーのみ
        shutil.copy2(input_path, output_path)
        print(f"  [SKIP]  非対応形式のためコピーのみ → {output_path.name}")


# ---------- メイン ----------

def main():
    parser = argparse.ArgumentParser(
        description="マスクツール: 機密情報を隠蔽して共有可能なファイルを生成する",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("input", help="入力ファイルまたはフォルダ")
    parser.add_argument("--config", required=True, help="マスク設定ファイル (YAML)")
    parser.add_argument("--output", help="出力先ファイル/フォルダ（省略時は自動命名）")
    parser.add_argument(
        "--suffix", default="_masked",
        help="出力ファイルに付加するサフィックス（デフォルト: _masked）"
    )
    args = parser.parse_args()

    rules = load_rules(args.config)
    if not rules:
        print("警告: マスクルールが0件です。mask_config.yaml を確認してください。")
        return

    print(f"マスクルール: {len(rules)}件 読み込み完了")

    input_path = Path(args.input)
    if not input_path.exists():
        sys.exit(f"エラー: 入力パスが見つかりません: {input_path}")

    if input_path.is_dir():
        output_dir = Path(args.output) if args.output else input_path.parent / (input_path.name + args.suffix)
        output_dir.mkdir(parents=True, exist_ok=True)
        files = sorted(f for f in input_path.rglob("*") if f.is_file())
        print(f"フォルダ処理: {len(files)}ファイル → {output_dir}/")
        for f in files:
            rel = f.relative_to(input_path)
            out = output_dir / rel
            out.parent.mkdir(parents=True, exist_ok=True)
            process_file(f, out, rules)
    else:
        if args.output:
            output_path = Path(args.output)
        else:
            output_path = input_path.parent / (input_path.stem + args.suffix + input_path.suffix)
        process_file(input_path, output_path, rules)

    print("完了")


if __name__ == "__main__":
    main()
